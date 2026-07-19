import json
import pytest

from knowledge_pipelines.documentation_engine import parsers
from knowledge_pipelines.documentation_engine.parsers import markdown, plaintext, pdf, docx, openapi


def test_markdown_extracts_heading_hierarchy(tmp_path):
    f = tmp_path / "doc.md"
    f.write_text("# Title\n\nSome intro text.\n\n## Section One\n\nBody text with a [link](http://example.com) and **bold**.\n\n### Subsection\n\nMore text.\n")
    result = markdown.parse(str(f))
    levels = [h["level"] for h in result["structure_metadata"]["headings"]]
    texts = [h["text"] for h in result["structure_metadata"]["headings"]]
    assert levels == [1, 2, 3]
    assert texts == ["Title", "Section One", "Subsection"]


def test_markdown_strips_syntax_from_clean_text(tmp_path):
    f = tmp_path / "doc.md"
    f.write_text("# Title\n\nA [link](http://example.com) and **bold** and `code`.\n")
    result = markdown.parse(str(f))
    assert "[link]" not in result["clean_text"]
    assert "**" not in result["clean_text"]
    assert "`" not in result["clean_text"]
    assert "link" in result["clean_text"]
    assert "bold" in result["clean_text"]
    assert "code" in result["clean_text"]


def test_plaintext_passes_through_unchanged(tmp_path):
    f = tmp_path / "notes.txt"
    f.write_text("Just some plain notes.\nSecond line.")
    result = plaintext.parse(str(f))
    assert "Just some plain notes." in result["clean_text"]
    assert result["structure_metadata"]["headings"] == []


def test_pdf_parses_a_real_generated_pdf_with_real_text(tmp_path):
    """
    Generates a genuine PDF file with real drawn text (via reportlab, not
    a mock or a blank page) and parses it back — a real round trip
    proving actual text extraction, not just page counting.
    """
    from reportlab.pdfgen import canvas

    f = tmp_path / "test.pdf"
    c = canvas.Canvas(str(f))
    c.drawString(72, 700, "Invoice Approval Policy")
    c.drawString(72, 680, "Invoices over five thousand dollars require manager sign-off.")
    c.showPage()
    c.save()

    result = pdf.parse(str(f))
    assert result["structure_metadata"]["format"] == "pdf"
    assert result["structure_metadata"]["page_count"] == 1
    assert result["structure_metadata"]["headings"][0]["text"] == "Page 1"
    assert "Invoice Approval Policy" in result["clean_text"]
    assert "manager sign-off" in result["clean_text"]


def test_docx_extracts_real_text_and_headings(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_heading("Architecture Overview", level=1)
    doc.add_paragraph("This document describes the system architecture.")
    doc.add_heading("Component Design", level=2)
    doc.add_paragraph("Details about components.")
    f = tmp_path / "arch.docx"
    doc.save(str(f))

    result = docx.parse(str(f))
    assert "This document describes the system architecture." in result["clean_text"]
    headings = result["structure_metadata"]["headings"]
    assert [h["text"] for h in headings] == ["Architecture Overview", "Component Design"]
    assert [h["level"] for h in headings] == [1, 2]


def test_openapi_extracts_endpoints_as_sections(tmp_path):
    spec = {
        "openapi": "3.0.0",
        "info": {"title": "Sample API", "version": "1.0"},
        "paths": {
            "/widgets": {
                "get": {"summary": "List widgets", "description": "Returns all widgets."},
                "post": {"summary": "Create a widget"},
            },
        },
    }
    f = tmp_path / "spec.json"
    f.write_text(json.dumps(spec))

    result = openapi.parse(str(f))
    assert "Sample API" in result["clean_text"]
    assert "List widgets" in result["clean_text"]
    assert "Returns all widgets." in result["clean_text"]
    headings = {h["text"] for h in result["structure_metadata"]["headings"]}
    assert "GET /widgets" in headings
    assert "POST /widgets" in headings


def test_openapi_parses_yaml_form(tmp_path):
    f = tmp_path / "spec.yaml"
    f.write_text("openapi: 3.0.0\ninfo:\n  title: YAML API\n  version: '2.0'\npaths:\n  /ping:\n    get:\n      summary: Health check\n")
    result = openapi.parse(str(f))
    assert "YAML API" in result["clean_text"]
    assert "Health check" in result["clean_text"]


def test_dispatch_picks_parser_by_extension(tmp_path):
    f = tmp_path / "readme.md"
    f.write_text("# Hello")
    assert parsers.parser_for(str(f)) is markdown


def test_dispatch_unsupported_extension_raises():
    with pytest.raises(parsers.UnsupportedFormat):
        parsers.parser_for("file.exe")


def test_dispatch_unparseable_document_raises_parse_failed(tmp_path):
    f = tmp_path / "broken.pdf"
    f.write_bytes(b"this is not a real pdf file")
    with pytest.raises(parsers.ParseFailed):
        parsers.parse(str(f))
